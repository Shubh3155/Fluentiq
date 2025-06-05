# vector_store_builder.py

import os
import io
from langchain.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class VectorDBCreator:
    def __init__(self):
        self.api_key = os.getenv("GOOGLE_API_KEY")
        if not self.api_key:
            raise ValueError("GOOGLE_API_KEY not found in environment variables")
        
        # Initialize embeddings
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=self.api_key
        )
        
        # Text splitter configuration
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", "! ", "? ", " "]
        )
    
    def extract_text_from_file(self, file_path):
        """Extract text from various file types"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            
            elif file_extension == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        return text
                except ImportError:
                    print("PyPDF2 not installed. Please install it: pip install PyPDF2")
                    return None
            
            elif file_extension == '.docx':
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except ImportError:
                    print("python-docx not installed. Please install it: pip install python-docx")
                    return None
            
            else:
                # Try to read as text file
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
                    
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            return None
    
    def create_vectorstore_from_text(self, text, vectorstore_name="faiss_vectorstore"):
        """Create FAISS vectorstore from text content"""
        try:
            # Create documents
            docs = [Document(page_content=text)]
            split_docs = self.text_splitter.split_documents(docs)
            
            if not split_docs:
                print("No documents created after splitting")
                return None
            
            print(f"Created {len(split_docs)} document chunks")
            
            # Create vectorstore
            vectorstore = FAISS.from_documents(split_docs, self.embeddings)
            
            # Save vectorstore to disk
            vectorstore.save_local(vectorstore_name)
            print(f"✅ Vectorstore saved as '{vectorstore_name}'")
            
            return vectorstore
            
        except Exception as e:
            print(f"Error creating vectorstore: {e}")
            return None
    
    def create_vectorstore_from_file(self, file_path, vectorstore_name="faiss_vectorstore"):
        """Create vectorstore from a file"""
        print(f"Processing file: {file_path}")
        
        text_content = self.extract_text_from_file(file_path)
        if not text_content:
            print("Failed to extract text from file")
            return None
        
        print(f"Extracted {len(text_content)} characters from file")
        return self.create_vectorstore_from_text(text_content, vectorstore_name)
    
    def create_vectorstore_from_multiple_files(self, file_paths, vectorstore_name="faiss_vectorstore"):
        """Create vectorstore from multiple files"""
        all_texts = []
        
        for file_path in file_paths:
            print(f"Processing file: {file_path}")
            text_content = self.extract_text_from_file(file_path)
            if text_content:
                all_texts.append(text_content)
                print(f"✅ Successfully processed {file_path}")
            else:
                print(f"❌ Failed to process {file_path}")
        
        if not all_texts:
            print("No text content extracted from any files")
            return None
        
        # Combine all texts
        combined_text = "\n\n".join(all_texts)
        print(f"Combined text length: {len(combined_text)} characters")
        
        return self.create_vectorstore_from_text(combined_text, vectorstore_name)
    
    def load_existing_vectorstore(self, vectorstore_name="faiss_vectorstore"):
        """Load existing vectorstore from disk"""
        try:
            vectorstore = FAISS.load_local(vectorstore_name, self.embeddings, allow_dangerous_deserialization=True)
            print(f"✅ Loaded existing vectorstore '{vectorstore_name}'")
            return vectorstore
        except Exception as e:
            print(f"❌ Could not load vectorstore '{vectorstore_name}': {e}")
            return None

def main():
    """Main function to demonstrate usage"""
    print("🚀 Vector Database Creator")
    print("=" * 50)
    
    try:
        # Initialize creator
        creator = VectorDBCreator()
        
        # Example usage - modify as needed
        print("\nChoose an option:")
        print("1. Create vectorstore from a single file")
        print("2. Create vectorstore from multiple files")
        print("3. Load existing vectorstore (test)")
        
        choice = input("\nEnter your choice (1-3): ").strip()
        
        if choice == "1":
            file_path = input("Enter file path: ").strip()
            vectorstore_name = input("Enter vectorstore name (default: faiss_vectorstore): ").strip()
            if not vectorstore_name:
                vectorstore_name = "faiss_vectorstore"
            
            if os.path.exists(file_path):
                vectorstore = creator.create_vectorstore_from_file(file_path, vectorstore_name)
                if vectorstore:
                    print(f"\n🎉 Successfully created vectorstore '{vectorstore_name}'!")
                    print(f"📁 Files created: {vectorstore_name}/ directory")
            else:
                print(f"❌ File not found: {file_path}")
        
        elif choice == "2":
            file_paths_input = input("Enter file paths (comma-separated): ").strip()
            file_paths = [path.strip() for path in file_paths_input.split(",")]
            vectorstore_name = input("Enter vectorstore name (default: faiss_vectorstore): ").strip()
            if not vectorstore_name:
                vectorstore_name = "faiss_vectorstore"
            
            # Check if all files exist
            valid_files = [fp for fp in file_paths if os.path.exists(fp)]
            if valid_files:
                vectorstore = creator.create_vectorstore_from_multiple_files(valid_files, vectorstore_name)
                if vectorstore:
                    print(f"\n🎉 Successfully created vectorstore '{vectorstore_name}'!")
                    print(f"📁 Files created: {vectorstore_name}/ directory")
            else:
                print("❌ No valid files found")
        
        elif choice == "3":
            vectorstore_name = input("Enter vectorstore name to load (default: faiss_vectorstore): ").strip()
            if not vectorstore_name:
                vectorstore_name = "faiss_vectorstore"
            
            vectorstore = creator.load_existing_vectorstore(vectorstore_name)
            if vectorstore:
                # Test search
                test_query = input("Enter a test query: ").strip()
                if test_query:
                    results = vectorstore.similarity_search(test_query, k=2)
                    print(f"\n🔍 Search results for '{test_query}':")
                    for i, doc in enumerate(results, 1):
                        print(f"\nResult {i}:")
                        print(doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content)
        
        else:
            print("❌ Invalid choice")
    
    except Exception as e:
        print(f"❌ Error: {e}")
    
    print("\n" + "=" * 50)
    print("💡 Tips:")
    print("- The vectorstore will be saved as a directory")
    print("- Copy this directory to your chatbot project")
    print("- Use the vectorstore name in your chatbot configuration")
    print("- Supported formats: .txt, .md, .pdf, .docx")

if __name__ == "__main__":
    main()